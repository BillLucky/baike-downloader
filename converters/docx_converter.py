"""
DOCX 转换器
使用 python-docx 按章节重建 Word 文档
"""
import os
from pathlib import Path
from datetime import datetime
from typing import Optional

from docx import Document
from docx.shared import Cm, Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

import sys
sys.path.insert(0, str(Path(__file__).parent.parent))
import config
from scraper.baike_parser import BaikeDocument


def set_run_font(run, east_asian: str, latin: str = "Times New Roman", size: int = 11):
    """设置 run 的字体"""
    run.font.east_asian = east_asian
    run.font.name = latin
    run.font.size = Pt(size)
    # 中文字体
    r = run._element
    rPr = r.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.insert(0, rFonts)
    rFonts.set(qn("w:eastAsia"), east_asian)


def add_heading(doc: Document, text: str, level: int = 1):
    """添加标题段落"""
    p = doc.add_heading(text, level=level)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    for run in p.runs:
        if level == 1:
            set_run_font(run, config.DOCX_CONFIG["heading1_font"],
                         "Arial", config.DOCX_CONFIG["heading1_size"])
            run.font.bold = True
            run.font.color.rgb = RGBColor(0x11, 0x11, 0x11)
        elif level == 2:
            set_run_font(run, config.DOCX_CONFIG["heading2_font"],
                         "Arial", config.DOCX_CONFIG["heading2_size"])
            run.font.bold = True
            run.font.color.rgb = RGBColor(0x22, 0x22, 0x22)
        else:
            set_run_font(run, config.DOCX_CONFIG["heading2_font"],
                         "Arial", 12)
            run.font.bold = True
    return p


def add_paragraph(doc: Document, text: str, indent: bool = False):
    """添加正文段落"""
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    p.paragraph_format.line_spacing = config.DOCX_CONFIG["line_spacing"]
    if indent:
        p.paragraph_format.first_line_indent = Cm(0.74)  # 2字符

    run = p.add_run(text)
    set_run_font(run, config.DOCX_CONFIG["body_font"],
                 config.DOCX_CONFIG["body_font"],
                 config.DOCX_CONFIG["body_size"])
    return p


def add_image(doc: Document, local_path: str, caption: str = ""):
    """添加图片（自动处理 WebP、SVG 等格式）"""
    try:
        from PIL import Image
        from io import BytesIO

        path = Path(local_path)
        if not path.exists():
            return

        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()

        ext = path.suffix.lower()
        # WebP / SVG → 转换为 PNG
        if ext in (".webp", ".svg"):
            try:
                if ext == ".svg":
                    import cairosvg
                    buf = BytesIO()
                    cairosvg.svg2png(bytestring=path.read_bytes(), output_buffer=buf)
                else:
                    img = Image.open(path)
                    buf = BytesIO()
                    img.save(buf, format="PNG")
                buf.seek(0)
                run.add_picture(buf, width=Cm(14))
            except Exception:
                # SVG/ WebP 转换失败，跳过
                p._element.getparent().remove(p._element)
                return
        else:
            with open(path, "rb") as f:
                run.add_picture(f, width=Cm(14))

        if caption:
            cap_p = doc.add_paragraph()
            cap_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            cap_run = cap_p.add_run(caption)
            set_run_font(cap_run, config.DOCX_CONFIG["body_font"],
                         config.DOCX_CONFIG["body_font"], 9)
            cap_run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)
            cap_run.font.italic = True
    except Exception as e:
        print(f"  ⚠ 图片添加失败: {local_path} -> {e}")


def add_table(doc: Document, items):
    """添加基本信息表格"""
    if not items:
        return
    table = doc.add_table(rows=len(items), cols=2)
    table.style = "Table Grid"
    for i, item in enumerate(items):
        row = table.rows[i]
        row.cells[0].text = item.label
        row.cells[1].text = item.value
        # 加粗标签列
        for run in row.cells[0].paragraphs[0].runs:
            run.font.bold = True
    doc.add_paragraph()  # 空行


class DOCXConverter:
    """DOCX 转换器"""

    def __init__(self):
        pass

    def convert(
        self,
        doc: BaikeDocument,
        output_path: Path,
        url_to_path: Optional[dict[str, str]] = None,
    ) -> Path:
        """
        将 BaikeDocument 转为 DOCX
        """
        url_to_path = url_to_path or {}
        document = Document()

        # 页面设置
        section = document.sections[0]
        section.page_width = Cm(config.DOCX_CONFIG["page_width"])
        section.page_height = Cm(config.DOCX_CONFIG["page_height"])
        section.top_margin = Cm(config.DOCX_CONFIG["margin_top"])
        section.bottom_margin = Cm(config.DOCX_CONFIG["margin_bottom"])
        section.left_margin = Cm(config.DOCX_CONFIG["margin_left"])
        section.right_margin = Cm(config.DOCX_CONFIG["margin_right"])

        # 标题
        add_heading(document, doc.title, level=1)

        # 副标题/摘要
        if doc.subtitle:
            p = add_paragraph(document, doc.subtitle)
            p.runs[0].font.color.rgb = RGBColor(0x66, 0x66, 0x66)

        # 元信息
        meta_p = document.add_paragraph()
        meta_run = meta_p.add_run(
            f"来源：{doc.url}\n"
            f"下载时间：{datetime.now().strftime('%Y-%m-%d %H:%M:%S')}"
        )
        meta_run.font.size = Pt(9)
        meta_run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
        meta_p.paragraph_format.space_after = Pt(12)

        # 分隔线
        p = document.add_paragraph()
        pPr = p._element.get_or_add_pPr()
        pBdr = OxmlElement("w:pBdr")
        bottom = OxmlElement("w:bottom")
        bottom.set(qn("w:val"), "single")
        bottom.set(qn("w:sz"), "6")
        bottom.set(qn("w:space"), "1")
        bottom.set(qn("w:color"), "CCCCCC")
        pBdr.append(bottom)
        pPr.append(pBdr)

        # 基本信息栏
        if doc.basic_info:
            add_heading(document, "基本信息", level=2)
            add_table(document, doc.basic_info)

        # 正文章节
        for section in doc.sections:
            add_heading(document, section.title, level=section.level)

            # 解析章节 HTML，处理文字和图片
            from bs4 import BeautifulSoup
            soup = BeautifulSoup(section.content_html or "", "lxml")

            for child in soup.children:
                if not hasattr(child, "name") or child.name is None:
                    continue
                tag_name = child.name.lower()

                if tag_name == "p":
                    text = child.get_text(strip=True)
                    if text:
                        add_paragraph(document, text, indent=True)

                elif tag_name in ("h2", "h3", "h4"):
                    # 已在上面用 add_heading 处理
                    pass

                elif tag_name == "img":
                    src = child.get("data-src") or child.get("src") or ""
                    if src and src in url_to_path:
                        alt = child.get("alt", "")
                        add_image(document, url_to_path[src], alt)

                elif tag_name in ("ul", "ol"):
                    for li in child.find_all("li", recursive=False):
                        text = li.get_text(strip=True)
                        if text:
                            p = document.add_paragraph(style="List Bullet")
                            run = p.add_run(text)
                            set_run_font(run, config.DOCX_CONFIG["body_font"],
                                         config.DOCX_CONFIG["body_font"],
                                         config.DOCX_CONFIG["body_size"])

                elif tag_name == "table":
                    rows = child.find_all("tr")
                    if rows:
                        tbl = document.add_table(rows=len(rows), cols=len(rows[0].find_all(["td", "th"])))
                        tbl.style = "Table Grid"
                        for ri, tr in enumerate(rows):
                            cells = tr.find_all(["td", "th"])
                            for ci, cell in enumerate(cells):
                                if ci < len(tbl.rows[ri].cells):
                                    tbl.rows[ri].cells[ci].text = cell.get_text(strip=True)

                elif tag_name == "div":
                    text = child.get_text(strip=True)
                    if text:
                        add_paragraph(document, text, indent=True)

        # 图片集（从 url_to_path 中取已下载图片）
        downloaded_images = [p for p in url_to_path.values() if Path(p).exists()]
        if downloaded_images:
            add_heading(document, "图片集", level=2)
            add_paragraph(document, f"（共 {len(downloaded_images)} 张图片）")
            for local_path in downloaded_images[:20]:  # 最多20张
                add_image(document, local_path, "")

        # 页脚
        for section in document.sections:
            footer = section.footer
            fp = footer.paragraphs[0]
            if not fp.runs:
                fp.add_run(f"来源：{doc.url}  ·  百度百科下载器")
            fp.runs[0].font.size = Pt(8)
            fp.runs[0].font.color.rgb = RGBColor(0xBB, 0xBB, 0xBB)

        document.save(str(output_path))
        print(f"  ✓ DOCX 已保存: {output_path.name}")
        return output_path
